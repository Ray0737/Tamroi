import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

FONT = 'TH SarabunPSK'
SZ_TITLE = 16    # pt  — cover / title pages
SZ_BODY  = 16    # pt  — body text
SZ_FORM  = 14    # pt  — form field rows
THAI_JUSTIFY = 9  # WD_ALIGN_PARAGRAPH value for Thai full justify

IMG_DIR = os.path.dirname(os.path.abspath(__file__))

def set_thai_justify(para):
    """Apply Thai full-justify alignment via raw XML (value 9)."""
    pPr = para._p.get_or_add_pPr()
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'thaiDistribute')
    pPr.append(jc)

def add_para(doc, text, size=SZ_BODY, bold=False, align=None,
             space_before=0, space_after=0, left_indent=None, thai_j=False):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    if left_indent is not None:
        pf.left_indent = Cm(left_indent)
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    # Force Thai font in rPr
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'),    FONT)
    rFonts.set(qn('w:hAnsi'),    FONT)
    rFonts.set(qn('w:cs'),       FONT)
    rFonts.set(qn('w:eastAsia'), FONT)
    rPr.insert(0, rFonts)
    if thai_j:
        set_thai_justify(p)
    return p

def add_heading(doc, text, size=SZ_BODY, space_before=6):
    return add_para(doc, text, size=size, bold=True,
                    space_before=space_before, space_after=2)

def add_body(doc, text, indent=0):
    return add_para(doc, text, size=SZ_BODY, space_after=0,
                    left_indent=indent if indent else None, thai_j=True)

def add_form_row(doc, label, value, indent=0):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after  = Pt(0)
    if indent:
        pf.left_indent = Cm(indent)
    for text, bold in [(label, False), ('\t', False), (value, False)]:
        run = p.add_run(text)
        run.font.name = FONT
        run.font.size = Pt(SZ_FORM)
        run.bold = bold
        rPr = run._r.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'),    FONT)
        rFonts.set(qn('w:hAnsi'),    FONT)
        rFonts.set(qn('w:cs'),       FONT)
        rFonts.set(qn('w:eastAsia'), FONT)
        rPr.insert(0, rFonts)
    return p

def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(docx_break_type())

def docx_break_type():
    from docx.enum.text import WD_BREAK
    return WD_BREAK.PAGE

def add_image(doc, filename, caption=None, width_cm=14):
    img_path = os.path.join(IMG_DIR, filename)
    if not os.path.exists(img_path):
        add_para(doc, f'[รูปภาพ: {filename}]', size=SZ_BODY,
                 align=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=2)
        if caption:
            add_para(doc, caption, size=SZ_FORM,
                     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after  = Pt(2)
    run = p.add_run()
    run.add_picture(img_path, width=Cm(width_cm))
    if caption:
        add_para(doc, caption, size=SZ_FORM,
                 align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)

def add_bullet(doc, text, indent=1.0):
    p = add_para(doc, f'- {text}', size=SZ_BODY,
                 left_indent=indent, space_after=0, thai_j=True)
    return p

def add_section_divider(doc):
    add_para(doc, '', size=SZ_BODY, space_before=4, space_after=4)

# ────────────────────────────────────────────
doc = Document()

# Page setup  — A4, 2.54 cm all margins
sec = doc.sections[0]
sec.page_width  = Cm(21.0)
sec.page_height = Cm(29.7)
sec.left_margin   = Cm(2.54)
sec.right_margin  = Cm(2.54)
sec.top_margin    = Cm(2.54)
sec.bottom_margin = Cm(2.54)

# ═══════════════════════════════════════════
# PAGE 1 — ข้อเสนอโครงการ (Cover form)
# ═══════════════════════════════════════════
add_para(doc, 'ชื่อโครงการ (ภาษาไทย) ตามรอย',
         size=SZ_TITLE, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc, '(ภาษาอังกฤษ) Tamroi',
         size=SZ_TITLE, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc, 'ประเภทโปรแกรมที่เสนอ โปรแกรมเพื่อการศึกษา (EDU Track)',
         size=SZ_TITLE, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc, 'ข้อเสนอโครงการ การแข่งขันพัฒนาโปรแกรมคอมพิวเตอร์แห่งประเทศไทย',
         size=SZ_TITLE, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc, 'รหัสโครงการ 26p23c0256',
         size=SZ_TITLE, bold=True)

add_section_divider(doc)
add_para(doc, 'ทีมพัฒนา หัวหน้าโครงการ', size=SZ_FORM, bold=True, space_before=4)
add_form_row(doc, '1. ชื่อ-นามสกุล', '\t\t\tนาย รพี รัตนมนูญพร (ชาย)')
add_form_row(doc, 'วัน/เดือน/ปีเกิด', '\t\t\t01 เมษายน 2553  14 ปี 6 เดือน')
add_form_row(doc, 'สถานศึกษาสาขา', '\t\tวิศวกรรมปัญญาประดิษฐ์ (AI)')
add_form_row(doc, '', '\t\t\tโรงเรียนสาธิต มหาวิทยาลัยศรีนครินทรวิโรฒ ประสานมิตร (ฝ่ายมัธยม)')
add_form_row(doc, 'ที่อยู่ตามทะเบียนบ้าน', '\t[ที่อยู่ตามทะเบียนบ้าน]')
add_form_row(doc, 'สถานที่ติดต่อ', '\t\tอาคารวิทยวิโรฒ เลขที่ 176 ซ.สุขุมวิท 23 แขวงคลองเตยเหนือ เขตวัฒนา กรุงเทพมหานคร 10110')
add_form_row(doc, 'โทรศัพท์', '\t\t\t\t02-260-9986')
add_form_row(doc, 'ระดับการศึกษา', '\t\t\tมัธยมศึกษาตอนปลาย')
add_form_row(doc, 'มือถือ', '\t\t\t\t0621132170')
add_form_row(doc, 'e-mail', '\t\t\t\traphee.rattanamanoonporn@gmail.com')
add_form_row(doc, 'ลงชื่อ', '\t\t\t\t')

add_section_divider(doc)
add_para(doc, 'ผู้ร่วมโครงการ', size=SZ_FORM, bold=True, space_before=4)
add_form_row(doc, '2. ชื่อ-นามสกุล', '\t\t\tนางสาว รชยา เชวงกิจวณิช (หญิง)')
add_form_row(doc, 'วัน/เดือน/ปีเกิด', '\t\t\t19 มิถุนายน 2551  16 ปี 4 เดือน')
add_form_row(doc, 'สถานศึกษา', '\t\t\tสาขา วิศวกรรมปัญญาประดิษฐ์ (AI)')
add_form_row(doc, '', '\t\t\t\tโรงเรียนสาธิต มหาวิทยาลัยศรีนครินทรวิโรฒ ประสานมิตร (ฝ่ายมัธยม)')
add_form_row(doc, 'ที่อยู่ตามทะเบียนบ้าน', '\t[ที่อยู่ตามทะเบียนบ้าน]')
add_form_row(doc, 'สถานที่ติดต่อ', '\t\t\tอาคารวิทยวิโรฒ เลขที่ 176 ซ.สุขุมวิท 23 แขวงคลองเตยเหนือ เขตวัฒนา กรุงเทพมหานคร 10110')
add_form_row(doc, 'โทรศัพท์', '\t\t\t\t02-260-9986')
add_form_row(doc, 'ระดับการศึกษา', '\t\t\tมัธยมศึกษาตอนปลาย')
add_form_row(doc, 'มือถือ', '\t\t\t\t0929309287')
add_form_row(doc, 'e-mail', '\t\t\t\tcharlotte.kamoshita00@gmail.com')
add_form_row(doc, 'ลงชื่อ', '\t\t\t\t')

add_section_divider(doc)
add_form_row(doc, '3. ชื่อ-นามสกุล', '\t\t\tนาย ปภาวิชญ์ แซ่หลิ่ว (ชาย)')
add_form_row(doc, 'วัน/เดือน/ปีเกิด', '\t\t\t02 สิงหาคม 2552  15 ปี 2 เดือน')
add_form_row(doc, 'สถานศึกษา', '\t\t\tสาขา วิศวกรรมปัญญาประดิษฐ์ (AI)')
add_form_row(doc, '', '\t\t\t\tโรงเรียนสาธิต มหาวิทยาลัยศรีนครินทรวิโรฒ ประสานมิตร (ฝ่ายมัธยม)')
add_form_row(doc, 'ที่อยู่ตามทะเบียนบ้าน', '\t[ที่อยู่ตามทะเบียนบ้าน]')
add_form_row(doc, 'สถานที่ติดต่อ', '\t\t\tอาคารวิทยวิโรฒ เลขที่ 176 ซ.สุขุมวิท 23 แขวงคลองเตยเหนือ เขตวัฒนา กรุงเทพมหานคร 10110')
add_form_row(doc, 'โทรศัพท์', '\t\t\t\t02-260-9986')
add_form_row(doc, 'ระดับการศึกษา', '\t\t\tมัธยมศึกษาตอนปลาย')
add_form_row(doc, 'มือถือ', '\t\t\t\t0917511470')
add_form_row(doc, 'e-mail', '\t\t\t\tpapawit@proton.me')
add_form_row(doc, 'ลงชื่อ', '\t\t\t\t')

add_section_divider(doc)
add_para(doc, 'อาจารย์ที่ปรึกษาโครงการ', size=SZ_FORM, bold=True, space_before=4)
add_form_row(doc, 'ชื่อ-นามสกุล', '\t\t\tนาย ธนภูมิ เรืองไพศาล (ชาย)')
add_form_row(doc, 'ระดับการศึกษา', '\t\t\tปริญญาตรี')
add_form_row(doc, 'ตำแหน่งทางวิชาการ', '\t\tครู/อาจารย์')
add_form_row(doc, 'สังกัด/สถาบัน', '\t\t\tสาขา วิศวกรรมปัญญาประดิษฐ์ โรงเรียนสาธิต มหาวิทยาลัยศรีนครินทรวิโรฒ ประสานมิตร (ฝ่ายมัธยม)')
add_form_row(doc, 'สถานที่ติดต่อ', '\t\t\tเลขที่ 176 แขวงคลองเตยเหนือ เขตวัฒนา กรุงเทพมหานคร 10110')
add_form_row(doc, 'โทรศัพท์', '\t\t\t\t02-260-9986')
add_form_row(doc, 'มือถือ', '\t\t\t\t082-156-3598')
add_form_row(doc, 'e-mail', '\t\t\t\ttanapoom.ru@spsm.ac.th')
add_body(doc, 'คำรับรอง "โครงการนี้เป็นความคิดริเริ่มของนักพัฒนาโครงการและไม่ได้ลอกเลียนแบบมาจากผู้อื่นผู้ใด ข้าพเจ้าขอรับรองว่าจะให้คำแนะนำและสนับสนุนให้นักพัฒนาในความดูแลของข้าพเจ้าดำเนินการศึกษา/วิจัย/พัฒนาตามหัวข้อที่เสนอและจะทำหน้าที่ประเมินผลงานดังกล่าวให้กับโครงการฯ ด้วย"')
add_form_row(doc, 'ลงชื่อ', '\t\t\t\t')

add_section_divider(doc)
add_para(doc, 'หัวหน้าสถาบัน (อธิการบดี/คณบดี/หัวหน้าภาควิชา/ผู้อำนวยการ/อาจารย์ใหญ่/หัวหน้าหมวด)',
         size=SZ_FORM, bold=True, space_before=4)
add_form_row(doc, 'ชื่อ-นามสกุล', '\t\t\tนาย กาลัญญู สูงใหญ่ (ชาย)')
add_form_row(doc, 'ตำแหน่งทางวิชาการ', '\t\tครู/อาจารย์')
add_form_row(doc, 'ตำแหน่งทางบริหาร', '\t\t\tหัวหน้า/ประธาน หมวด/สาขาวิชา')
add_form_row(doc, 'สังกัด/สถาบัน', '\t\t\tสาขา วิศวกรรมปัญญาประดิษฐ์ (AI) โรงเรียนสาธิต มหาวิทยาลัยศรีนครินทรวิโรฒ ประสานมิตร (ฝ่ายมัธยม)')
add_form_row(doc, 'สถานที่ติดต่อ', '\t\t\tอาคารวิทยวิโรฒ เลขที่ 176 ซ.สุขุมวิท 23 แขวงคลองเตยเหนือ เขตวัฒนา กรุงเทพมหานคร 10110')
add_form_row(doc, 'โทรศัพท์', '\t\t\t\t02-260-9986')
add_form_row(doc, 'มือถือ', '\t\t\t\t-')
add_body(doc, 'คำรับรอง "ข้าพเจ้าขอรับรองว่าผู้พัฒนามีสิทธิ์ขอรับทุนสนับสนุนตามเงื่อนไขที่โครงการฯกำหนดและอนุญาต ให้ดำเนินการศึกษา/วิจัย/พัฒนาตามหัวข้อที่ได้เสนอมานี้ในสถาบันได้ภายใต้การบังคับบัญชาของข้าพเจ้า"')
add_form_row(doc, 'e-mail', '\t\t\t\t')
add_form_row(doc, 'ลงชื่อ', '\t\t\t\t')

# ═══════════════════════════════════════════
# PAGE 2 — รายงานฉบับสมบูรณ์ (Title Page)
# ═══════════════════════════════════════════
doc.add_page_break()

for text in ['รายงานฉบับสมบูรณ์', '', 'ตามรอย', 'Tamroi', '',
             'โปรแกรมเพื่อการศึกษา ( นักเรียน )', '', '', '',
             'เสนอต่อ',
             'ศูนย์เทคโนโลยีอิเล็กทรอนิกส์และคอมพิวเตอร์แห่งชาติ',
             'สำนักงานพัฒนาวิทยาศาสตร์และเทคโนโลยีแห่งชาติ',
             'กระทรวงวิทยาศาสตร์และเทคโนโลยี', '', '',
             'ได้รับทุนอุดหนุนโครงการวิจัยพัฒนาและวิศวกรรม',
             'โครงการแข่งขันพัฒนาโปรแกรมคอมพิวเตอร์แห่งประเทศไทยครั้งที่ 26 ประจำปีงบประมาณ 2567',
             '', '', 'โดย', '',
             'นาย รพี รัตนมนูญพร',
             'นางสาว รชยา เชวงกิจวณิช',
             'นาย ปภาวิชญ์ แซ่หลิ่ว',
             '', 'อาจารย์ที่ปรึกษา', 'อาจารย์ธนภูมิ เรืองไพศาล',
             '', 'โรงเรียนสาธิต มหาวิทยาลัยศรีนครินทรวิโรฒ ประสานมิตร (ฝ่ายมัธยม)']:
    add_para(doc, text, size=SZ_TITLE, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)

# ═══════════════════════════════════════════
# กิตติกรรมประกาศ
# ═══════════════════════════════════════════
doc.add_page_break()
add_para(doc, 'กิตติกรรมประกาศ', size=SZ_TITLE, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
add_body(doc, 'คณะผู้จัดทำโครงการ "ตามรอย" ขอกราบขอบพระคุณผู้สนับสนุนที่ทำให้โปรแกรมถูกพัฒนา ขอขอบคุณโครงการแข่งขันพัฒนาโปรแกรมคอมพิวเตอร์แห่งประเทศไทยครั้งที่ 26 ที่เปิดโอกาสในการเข้าร่วมการแข่งขันพัฒนาโปรแกรมคอมพิวเตอร์และมอบทุนการพัฒนาทำให้ผู้จัดทำได้มีโอกาสพัฒนาตนเองและแสดงความรู้ความสามารถได้อย่างเต็มที่')
add_body(doc, 'ขอขอบคุณโรงเรียนสาธิต มหาวิทยาลัยศรีนครินทรวิโรฒ ประสานมิตร (ฝ่ายมัธยม) และ อาจารย์ ธนภูมิ เรืองไพศาล ที่ให้การสนับสนุนและให้คำปรึกษาเกี่ยวกับการพัฒนาโปรแกรมในการแข่งขันและเอกสารในการแข่ง ที่ทำให้โครงงานสำเร็จลุล่วง')
add_section_divider(doc)
add_para(doc, 'ผู้พัฒนา', size=SZ_TITLE, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER)

# ═══════════════════════════════════════════
# บทคัดย่อ
# ═══════════════════════════════════════════
doc.add_page_break()
add_para(doc, 'บทคัดย่อ', size=SZ_TITLE, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
add_body(doc,
    'โครงงานพัฒนาซอฟต์แวร์ ประเภทโปรแกรมเพื่อการศึกษา "ตามรอย" เป็นเว็บแอปพลิเคชันเพื่อการเรียนรู้ประวัติศาสตร์ไทยผ่านการสำรวจพื้นที่จริงในรูปแบบ Gamification โดยมุ่งเน้นการแก้ปัญหาการขาดแรงจูงใจในการเรียนประวัติศาสตร์ของเยาวชนไทย ผ่านกระบวนการเรียนรู้เชิงประสบการณ์ (Experiential Learning) ที่เชื่อมโยงสถานที่จริงเข้ากับองค์ความรู้ทางประวัติศาสตร์')
add_body(doc,
    'เว็บแอปพลิเคชัน "ตามรอย" ออกแบบตามหลักการ Situated Learning ซึ่งเชื่อว่าการเรียนรู้จะมีประสิทธิผลสูงสุดเมื่อเกิดขึ้นในบริบทที่แท้จริง ผู้ใช้งานจะเริ่มต้นด้วยแผนที่ประเทศไทยที่ถูกปกคลุมด้วย "หมอกแห่งสงคราม" (Fog of War) ซึ่งจะจางหายไปเมื่อผู้ใช้เดินทางไปยังพื้นที่จริงและผ่านการทดสอบความรู้ทางประวัติศาสตร์ที่เชื่อมโยงกับพื้นที่นั้นๆ กระบวนการนี้ส่งเสริมการเรียนรู้แบบ Active Learning เนื่องจากผู้เรียนต้องมีส่วนร่วมทั้งทางกายภาพ (การเดินทาง) และทางปัญญา (การตอบคำถามเนื้อหาประวัติศาสตร์) ระบบมีฟีเจอร์หลักดังนี้')
for item in ['1. หน้า Login / Sign-in (Google OAuth / Email)',
             '2. หน้า Map หลัก พร้อม Fog of War overlay',
             '3. หน้า Capture ตัวละครประวัติศาสตร์ (Quiz Modal พร้อมเนื้อหาประวัติศาสตร์)',
             '4. หน้า Archive (Historical Figures Collection พร้อมชีวประวัติย่อ)',
             '5. หน้า Leaderboard',
             '6. หน้า Profile Dashboard']:
    add_para(doc, item, size=SZ_BODY, left_indent=1.0, space_after=0)
add_body(doc,
    'โดยใช้ HTML5, CSS3, JavaScript (Vanilla ES6+) และ Leaflet.js ในการพัฒนาฝั่ง Frontend และ Supabase (PostgreSQL) ในการจัดการฐานข้อมูลและระบบ Authentication')
add_section_divider(doc)
add_body(doc, 'คำสำคัญ : การเรียนรู้ประวัติศาสตร์, Gamification, Situated Learning, Active Learning, Fog of War, Interactive Map, ประวัติศาสตร์ไทย, บุคคลสำคัญ')

# ═══════════════════════════════════════════
# Abstract
# ═══════════════════════════════════════════
doc.add_page_break()
add_para(doc, 'Abstract', size=SZ_TITLE, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
add_body(doc,
    '"Tamroi" (ตามรอย, meaning "Trace the Footsteps") is an educational web application designed to address the declining interest in Thai history among young learners. The project applies the principles of Situated Learning and Experiential Learning — delivering historical knowledge at the exact physical locations where history occurred, rather than through passive textbook study.')
add_body(doc,
    'Users are presented with a map of Thailand covered in Fog of War. As they physically travel to and check in at real locations, the fog lifts and the application prompts them with a history quiz tied to that area. Successfully answering unlocks a historical figure card (e.g., King Naresuan the Great at Ayutthaya, Sunthon Phu in Bangkok) with a biographical summary, reinforcing content through immediate contextual reward. This loop — travel, discover, learn, collect — follows the Active Learning cycle, ensuring that knowledge is acquired through engagement rather than memorization.')
add_body(doc, 'Key features:')
for item in ['1. Login / Sign-in page (Google OAuth / Email)',
             '2. Main Map page with Fog of War overlay',
             '3. Historical Figure Capture page (history quiz with biographical content)',
             '4. Archive page (Historical Figures Collection with short biographies)',
             '5. Leaderboard page',
             '6. Profile Dashboard page']:
    add_para(doc, item, size=SZ_BODY, left_indent=1.0, space_after=0)
add_body(doc,
    'Developed using HTML5, CSS3, JavaScript (Vanilla ES6+), Leaflet.js, and Supabase (PostgreSQL), Tamroi aims to transform Thai history education from a passive classroom activity into an active, location-based learning experience accessible to students and the general public.')
add_section_divider(doc)
add_body(doc, 'Keywords: Thai History Education, Gamification, Situated Learning, Active Learning, Experiential Learning, Fog of War, Interactive Map, Historical Figures')

# ═══════════════════════════════════════════
# สารบัญ
# ═══════════════════════════════════════════
doc.add_page_break()
add_para(doc, 'สารบัญ', size=SZ_TITLE, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
toc_items = [
    ('บทคัดย่อ', 'ก'), ('Abstract', 'ข'), ('สาระสำคัญของโครงงาน', 'ค'),
    ('ความสำคัญและความเป็นมาของโครงการ', 'ง'), ('สารบัญ', 'จ'),
    ('วัตถุประสงค์', '1'), ('ทฤษฎีที่เกี่ยวข้อง', '1'),
    ('เครื่องมือที่ใช้พัฒนา', '3'), ('ขอบเขตการใช้งาน', '4'),
    ('คุณลักษณะของอุปกรณ์ที่ใช้กับเว็บแอปพลิเคชันที่พัฒนา', '4'),
    ('หลักการทำงาน', '4'), ('กลุ่มผู้ใช้', '4'),
    ('โครงร่างของโปรแกรมพอสังเขป', '5'),
    ('เอกสารอ้างอิง', '11'), ('ภาคผนวก / ทีมพัฒนา', '12'),
]
for title, page in toc_items:
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after  = Pt(0)
    tab_stops = pf.tab_stops
    from docx.shared import Cm as CM
    from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
    tab_stops.add_tab_stop(CM(14), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    r1 = p.add_run(title)
    r1.font.name = FONT; r1.font.size = Pt(SZ_BODY)
    r2 = p.add_run(f'\t{page}')
    r2.font.name = FONT; r2.font.size = Pt(SZ_BODY)

# ═══════════════════════════════════════════
# บทนำ
# ═══════════════════════════════════════════
doc.add_page_break()
add_heading(doc, 'บทนำ (ความสำคัญและความเป็นมาของโครงการ)', size=SZ_TITLE, space_before=0)
add_body(doc,
    'จากผลสำรวจของสถาบันทดสอบทางการศึกษาแห่งชาติ (สทศ.) พบว่าวิชาประวัติศาสตร์เป็นหนึ่งในวิชาที่นักเรียนไทยทำคะแนนเฉลี่ยได้ต่ำที่สุดในการสอบระดับชาติ สอดคล้องกับงานวิจัยที่พบว่าเยาวชนไทยส่วนใหญ่มองว่าประวัติศาสตร์เป็นวิชาที่น่าเบื่อ ยากจดจำ และขาดความเชื่อมโยงกับชีวิตประจำวัน (กรมศิลปากร, 2566) ปัญหานี้มีรากมาจากรูปแบบการเรียนการสอนที่เน้นการท่องจำข้อมูลจากตำราเรียนเพียงอย่างเดียว โดยไม่ได้สร้างประสบการณ์ตรงหรือบริบทที่ทำให้ผู้เรียนรู้สึกว่าประวัติศาสตร์มีความหมายและมีชีวิตจริงๆ')
add_body(doc,
    'ในทางทฤษฎีการศึกษา Lave & Wenger (1991) ได้เสนอแนวคิด Situated Learning ซึ่งระบุว่าการเรียนรู้จะมีประสิทธิผลสูงสุดเมื่อเกิดขึ้นในบริบทและสภาพแวดล้อมที่ความรู้นั้นถูกนำไปใช้จริง กล่าวคือ การเรียนรู้เรื่องพระนเรศวรมหาราชที่จังหวัดพระนครศรีอยุธยาจะมีความหมายและติดทนกว่าการอ่านเรื่องราวเดียวกันในห้องเรียนอย่างมีนัยสำคัญ ประกอบกับแนวคิดของ Kolb (1984) เกี่ยวกับ Experiential Learning Cycle ที่เน้นกระบวนการ "ประสบการณ์ → สะท้อน → คิดวิเคราะห์ → ลงมือทำ" ซึ่งตรงกับกระบวนการที่ "ตามรอย" ออกแบบไว้')
add_body(doc,
    'นอกจากนี้ Deterding et al. (2011) ได้แสดงให้เห็นว่า Gamification สามารถเพิ่ม Engagement และแรงจูงใจในการเรียนรู้ได้อย่างมีนัยสำคัญ โดยเฉพาะในกลุ่มผู้เรียนที่เป็นเยาวชนซึ่งคุ้นเคยกับสื่อดิจิทัลและวิดีโอเกมอยู่แล้ว กลไกอย่างระบบสะสม (Collection) การจัดอันดับ (Leaderboard) และการปลดล็อค (Unlocking) ล้วนกระตุ้นแรงจูงใจภายใน (Intrinsic Motivation) ของผู้เรียน')
add_body(doc,
    'จากสถานการณ์และหลักการข้างต้น คณะผู้พัฒนาจึงออกแบบ "ตามรอย" ในฐานะเครื่องมือทางการศึกษา (Educational Tool) ที่ใช้เทคโนโลยีเว็บและระบบภูมิสารสนเทศ (Geospatial Technology) เพื่อสร้างประสบการณ์การเรียนประวัติศาสตร์ไทยแบบ Situated และ Active โดยผู้เรียนจะ "ตามรอย" บุคคลสำคัญในประวัติศาสตร์ไทยผ่านการเดินทางไปยังพื้นที่จริง ซึ่งถือเป็นการสืบสานมรดกทางวัฒนธรรมในรูปแบบที่สอดคล้องกับวิถีชีวิตของคนรุ่นใหม่')

# ═══════════════════════════════════════════
# วัตถุประสงค์
# ═══════════════════════════════════════════
add_section_divider(doc)
add_heading(doc, 'วัตถุประสงค์ของโครงงาน', size=SZ_TITLE)
for item in [
    'เพื่อพัฒนาสื่อการเรียนรู้ประวัติศาสตร์ไทยในรูปแบบ Gamification ที่ส่งเสริมการเรียนรู้เชิงประสบการณ์ (Experiential Learning) ผ่านการสำรวจพื้นที่จริง',
    'เพื่อเพิ่มแรงจูงใจและ Engagement ในการเรียนรู้ประวัติศาสตร์และบุคคลสำคัญของไทยในกลุ่มเยาวชน',
    'เพื่อสร้างฐานข้อมูลองค์ความรู้ทางประวัติศาสตร์ที่เชื่อมโยงกับพิกัดภูมิศาสตร์จริง สำหรับใช้เป็นทรัพยากรการเรียนรู้นอกห้องเรียน (Out-of-classroom Learning Resource)',
]:
    add_bullet(doc, item)

# ═══════════════════════════════════════════
# เป้าหมายและขอบเขต
# ═══════════════════════════════════════════
add_section_divider(doc)
add_heading(doc, 'เป้าหมายและขอบเขตของโครงการ', size=SZ_TITLE)
add_heading(doc, 'กลุ่มผู้ใช้โปรแกรม', size=SZ_BODY)
for item in [
    'กลุ่มหลัก: นักเรียนระดับมัธยมศึกษาและอุดมศึกษาที่ต้องการเรียนรู้ประวัติศาสตร์ไทยนอกห้องเรียน',
    'กลุ่มรอง: ครูและผู้ปกครองที่ต้องการสื่อเสริมการเรียนรู้ภาคสนาม (Field-based Learning)',
    'กลุ่มทั่วไป: นักท่องเที่ยวและผู้สนใจประวัติศาสตร์ที่ต้องการแรงจูงใจในการสำรวจสถานที่จริง',
]:
    add_bullet(doc, item)
add_heading(doc, 'การทำงานของเว็บแอปพลิเคชัน แบ่งออกเป็น 2 ส่วน ดังนี้', size=SZ_BODY)
add_bullet(doc, 'ส่วนที่ 1  ระบบแผนที่และการสำรวจภาคสนาม (Field Exploration System) — กระตุ้นการเดินทางไปยังสถานที่จริง')
add_bullet(doc, 'ส่วนที่ 2  ระบบเนื้อหาประวัติศาสตร์และการประเมินผลการเรียนรู้ (Historical Content & Learning Assessment) — Quiz, Archive และ Profile')

# ═══════════════════════════════════════════
# รายละเอียดการพัฒนา — เทคโนโลยี
# ═══════════════════════════════════════════
add_section_divider(doc)
add_heading(doc, 'รายละเอียดของการพัฒนา', size=SZ_TITLE)
add_heading(doc, 'ทฤษฎีหลักการและเทคโนโลยีที่ใช้', size=SZ_BODY)

add_heading(doc, '1.1. JavaScript (Vanilla ES6+)', size=SZ_BODY)
add_body(doc, 'JavaScript เป็นภาษาสคริปต์ที่ทำงานบนเบราว์เซอร์โดยตรง ไม่ต้องพึ่งพาเฟรมเวิร์กภายนอก การใช้ Vanilla ES6+ ทำให้สามารถควบคุม Logic ของแอปพลิเคชันได้อย่างยืดหยุ่น รองรับ Async/Await สำหรับการติดต่อ API, Module System สำหรับการจัดการโค้ด และ ES6 Class สำหรับโครงสร้างข้อมูลภายในแอปพลิเคชัน')

add_heading(doc, '1.2. Leaflet.js', size=SZ_BODY)
add_body(doc, 'Leaflet.js เป็น JavaScript Library โอเพ่นซอร์สสำหรับการแสดงผลแผนที่แบบ Interactive น้ำหนักเบา (39KB gzipped) รองรับทั้ง Desktop และ Mobile รองรับการเพิ่ม Custom Layer, Polygon Overlay และ Canvas Rendering ซึ่งเป็นพื้นฐานสำคัญของระบบ Fog of War ในโครงการนี้ โดยสามารถแสดงผลข้อมูล GeoJSON ของขอบเขตจังหวัดและเขตของประเทศไทยได้อย่างมีประสิทธิภาพ')

add_heading(doc, '1.3. Supabase (PostgreSQL)', size=SZ_BODY)
add_body(doc, 'Supabase เป็นแพลตฟอร์ม Backend-as-a-Service โอเพ่นซอร์สที่สร้างบน PostgreSQL มีความสามารถด้าน Database (Tables, Views, Functions, Triggers), Authentication (Email/Password, Google OAuth 2.0) และ Realtime Subscriptions ในโครงการนี้ใช้ Supabase เพื่อจัดเก็บข้อมูลผู้ใช้ ข้อมูลการสำรวจ และระบบ Ranking รวมถึงใช้ Database Triggers เพื่ออัปเดตคะแนน Legacy Score โดยอัตโนมัติ')

add_heading(doc, '1.4 เทคโนโลยีภูมิสารสนเทศ (Geospatial Technology)', size=SZ_BODY)
add_body(doc, 'Geolocation API เป็น Web API มาตรฐานที่รองรับโดยเบราว์เซอร์ทุกตัว ใช้สำหรับระบุตำแหน่งพิกัด GPS ของผู้ใช้แบบ Real-time ร่วมกับ GeoJSON ซึ่งเป็นรูปแบบมาตรฐานสำหรับข้อมูลทางภูมิศาสตร์ เพื่อทำ Polygon Matching ตรวจสอบว่าผู้ใช้อยู่ในเขตใดของประเทศไทย')

# ═══════════════════════════════════════════
# ทฤษฎีการเรียนรู้
# ═══════════════════════════════════════════
add_section_divider(doc)
add_heading(doc, 'ทฤษฎีการเรียนรู้และ Gamification เพื่อการศึกษาประวัติศาสตร์', size=SZ_BODY)

add_heading(doc, '2.1 Situated Learning (การเรียนรู้เชิงบริบท)', size=SZ_BODY)
add_body(doc, 'Lave & Wenger (1991) เสนอว่าการเรียนรู้ที่มีความหมายต้องเกิดขึ้น "ในบริบท" (in context) ไม่ใช่แบบ Abstract ในห้องเรียน "ตามรอย" นำหลักการนี้มาใช้โดยตรง — ผู้เรียนไม่ได้อ่านว่าสุนทรภู่เคยอาศัยอยู่ที่ไหน แต่ต้องเดินทางไปยังพื้นที่นั้นจริง ซึ่งทำให้ความรู้ถูกเข้ารหัสในความทรงจำระยะยาว (Long-term Memory Encoding) ได้ดีกว่าการท่องจำจากตำรา')

add_heading(doc, '2.2 Experiential Learning Cycle (วงจรการเรียนรู้จากประสบการณ์)', size=SZ_BODY)
add_body(doc, 'Kolb (1984) อธิบายวงจรการเรียนรู้ 4 ขั้น ซึ่ง "ตามรอย" รองรับครบทุกขั้น ดังนี้')
for row in [
    ('Concrete Experience (ประสบการณ์ตรง)', 'การเดินทางไปยังสถานที่ประวัติศาสตร์จริง'),
    ('Reflective Observation (สังเกตและสะท้อน)', 'การอ่านเนื้อหาประวัติศาสตร์ของพื้นที่นั้น'),
    ('Abstract Conceptualization (สร้างความเข้าใจ)', 'การตอบ Quiz เพื่อทดสอบความเข้าใจ'),
    ('Active Experimentation (ลงมือทดลอง)', 'การ Capture ตัวละครและสะสมความรู้ใหม่ต่อไป'),
]:
    add_para(doc, f'  {row[0]}  →  {row[1]}', size=SZ_BODY, left_indent=1.0, space_after=0)

add_heading(doc, '2.3 Gamification ในบริบทการศึกษาประวัติศาสตร์', size=SZ_BODY)
add_body(doc, '2.3.1 การสะสมและรางวัล (Collection & Reward) — ระบบการ Capture ตัวละครประวัติศาสตร์พร้อมชีวประวัติย่อกระตุ้นให้ผู้เรียนอยากสะสมความรู้เกี่ยวกับบุคคลสำคัญในยุคสมัยต่างๆ ระบบ Rarity Tier (S/A/C) สะท้อนความสำคัญทางประวัติศาสตร์ของบุคคล')
add_body(doc, '2.3.2 การแข่งขันทางสังคม (Social Competition) — ระบบ Leaderboard ที่วัดจาก Legacy Score และ Map Discovery Percentage กระตุ้นให้ผู้เรียนอยากสำรวจพื้นที่ใหม่และเรียนรู้เนื้อหาใหม่อย่างต่อเนื่อง')
add_body(doc, '2.3.3 การสำรวจและค้นพบ (Exploration & Discovery) — กลไก Fog of War กระตุ้น Curiosity Gap สมองมนุษย์รู้สึกไม่สบายเมื่อเห็นข้อมูลที่ "ซ่อนอยู่" และมีแรงจูงใจที่จะเปิดเผยมัน ในกรณีนี้สิ่งที่ซ่อนอยู่คือเรื่องราวทางประวัติศาสตร์ที่รอการค้นพบ')
add_body(doc, '2.3.4 การเล่าเรื่องเชิงประวัติศาสตร์ (Historical Narrative) — บุคคลทุกคนในระบบเชื่อมโยงกับพื้นที่จริง เช่น สมเด็จพระนเรศวรมหาราช ↔ พระนครศรีอยุธยา, สุนทรภู่ ↔ กรุงเทพ (บางกอก) สร้าง Mental Map ของประวัติศาสตร์ไทยในหัวผู้เรียน')

# ═══════════════════════════════════════════
# คุณสมบัติสำคัญ
# ═══════════════════════════════════════════
add_section_divider(doc)
add_heading(doc, 'คุณสมบัติสำคัญของเว็บแอปพลิเคชัน', size=SZ_BODY)
add_heading(doc, '3.1 ขอบเขตของโปรแกรม', size=SZ_BODY)
add_body(doc, 'เว็บแอปพลิเคชันที่รองรับการใช้งานผ่านเบราว์เซอร์บน Desktop และ Mobile โดยให้บริการครอบคลุมพื้นที่ประเทศไทยทั้ง 77 จังหวัด รองรับภาษาไทยเป็นหลัก')
add_heading(doc, '3.2 ลักษณะของอุปกรณ์ที่ใช้', size=SZ_BODY)
add_body(doc, 'คอมพิวเตอร์, Smartphone หรือ Tablet ที่มีเบราว์เซอร์รุ่นใหม่ (Chrome, Firefox, Safari, Edge) และรองรับ Geolocation API')
add_heading(doc, 'ภาษาและเครื่องมือที่ใช้ในการพัฒนา', size=SZ_BODY)
for item in ['HTML5, CSS3, JavaScript (Vanilla ES6+) สำหรับ Frontend',
             'Leaflet.js สำหรับ Interactive Map และ Fog of War',
             'Supabase (PostgreSQL) สำหรับ Backend, Database และ Authentication',
             'Figma สำหรับออกแบบ UI Mockup',
             'Git / GitHub สำหรับ Version Control และ Deployment']:
    add_bullet(doc, item)

add_heading(doc, 'หลักการทำงาน', size=SZ_BODY)
add_body(doc, 'ผู้ใช้เข้าสู่ระบบด้วย Google Account หรือ Email จากนั้นพบกับแผนที่ประเทศไทยที่ถูกปกคลุมด้วยหมอก (Fog of War) เมื่อผู้ใช้ทำการ Check-in ในเขตใดเขตหนึ่งผ่านระบบ GPS ระบบจะทำการ Polygon Matching เพื่อระบุเขตที่ผู้ใช้อยู่ แล้วลบหมอกออกจากแผนที่ในเขตนั้น จากนั้นระบบจะเปิด Quiz สำหรับ "Capture" ตัวละครประวัติศาสตร์ระดับ S, A หรือ C ที่เชื่อมโยงกับพื้นที่ คะแนน Legacy Score จะสะสมและแสดงผลบน Leaderboard โดยอัตโนมัติผ่าน Database Trigger ของ Supabase')

# ═══════════════════════════════════════════
# Input - Output Specification (with UI screenshots)
# ═══════════════════════════════════════════
doc.add_page_break()
add_heading(doc, 'Input - Output Specification', size=SZ_TITLE)
add_heading(doc, 'ตัวอย่างเว็บแอปพลิเคชัน', size=SZ_BODY)

add_heading(doc, '1.1) หน้า Login / Sign-in', size=SZ_BODY)
add_body(doc, 'ผู้ใช้เข้าสู่ระบบได้ 2 ช่องทาง ได้แก่ กด "Continue with Google" เชื่อมต่อ Google OAuth 2.0 เข้าสู่ระบบทันที หรือกรอก Email / Password เพื่อสมัครสมาชิกหรือเข้าสู่ระบบ')
add_image(doc, 'ui-login.png', 'รูปที่ 1 หน้า Login')
add_image(doc, 'ui-register.png', 'รูปที่ 2 หน้า Sign-up / Register')

add_heading(doc, '1.2) หน้า Map หลัก (Fog of War Map)', size=SZ_BODY)
add_body(doc, 'แสดงแผนที่ประเทศไทยที่ถูกปกคลุมด้วยหมอก (Fog of War) สีเข้ม หมอกจะจางหายไปทีละเขต/จังหวัดเมื่อผู้ใช้ Check-in ในพื้นที่จริง แสดงผลเป็น Percentage ของพื้นที่ที่ค้นพบแล้ว')
add_image(doc, 'ui-map.png', 'รูปที่ 3 หน้าแผนที่หลัก (Fog of War)')
add_image(doc, 'ui-home.png', 'รูปที่ 4 หน้า Home / Dashboard')

add_heading(doc, '1.3) หน้า Mission / Outpost (Check-in และ Capture ตัวละคร)', size=SZ_BODY)
add_body(doc, 'เมื่อผู้ใช้เข้าถึง Outpost ในพื้นที่ ระบบจะตรวจสอบพิกัด GPS ผ่าน Geolocation API แล้วเปิดภารกิจ (Mission) สำหรับ Capture ตัวละครประวัติศาสตร์ที่เชื่อมโยงกับพื้นที่นั้น เช่น อยุธยา → สมเด็จพระนเรศวรมหาราช (S-Tier)')
add_image(doc, 'ui-outpost.png', 'รูปที่ 5 หน้า Outpost')
add_image(doc, 'ui-mission.png', 'รูปที่ 6 หน้า Mission / Quiz ประวัติศาสตร์')

add_heading(doc, '1.4) หน้า Watchtower (บุคคลสำคัญทางประวัติศาสตร์)', size=SZ_BODY)
add_body(doc, 'แสดงรายละเอียดของบุคคลสำคัญทางประวัติศาสตร์ที่ค้นพบในพื้นที่ พร้อมชีวประวัติย่อและบริบทเชิงประวัติศาสตร์ที่เชื่อมโยงกับสถานที่นั้น')
add_image(doc, 'ui-watchtower-done.png', 'รูปที่ 7 Watchtower — ค้นพบแล้ว')
add_image(doc, 'ui-watchtower-uncomplete.png', 'รูปที่ 8 Watchtower — ยังไม่ค้นพบ')

add_heading(doc, '1.5) หน้า Archive (Historical Figures Collection)', size=SZ_BODY)
add_body(doc, 'แสดงรายการตัวละครประวัติศาสตร์ทั้งหมดที่สะสมได้ แบ่งตาม Rarity Tier (S / A / C) พร้อมข้อมูลชีวประวัติย่อและยุคสมัย เพื่อให้ผู้เรียนสามารถทบทวนความรู้ทางประวัติศาสตร์ที่ได้เรียนรู้ไปแล้วได้')
add_image(doc, 'ui-collections.png', 'รูปที่ 9 หน้า Archive / Collections')

add_heading(doc, '1.6) หน้า Leaderboard (การจัดอันดับ)', size=SZ_BODY)
add_body(doc, 'แสดงอันดับ Legacy Score ของผู้เรียนทั้งหมด เรียงจากคะแนนสูงสุด พร้อมข้อมูล Map Discovery Percentage และจำนวนตัวละครที่สะสมได้ เพื่อกระตุ้นการแข่งขันและแรงจูงใจในการเรียนรู้ต่อไป')
add_image(doc, 'ui-ranking1.png', 'รูปที่ 10 หน้า Leaderboard (1)')
add_image(doc, 'ui-ranking2.png', 'รูปที่ 11 หน้า Leaderboard (2)')

add_heading(doc, '1.7) หน้า Notification และ Settings', size=SZ_BODY)
add_body(doc, 'ระบบแจ้งเตือนเมื่อมีเหตุการณ์สำคัญ เช่น มี Outpost ใหม่ในพื้นที่ใกล้เคียง หรือมีการอัปเดต Leaderboard รวมถึงหน้า Settings สำหรับจัดการบัญชีและการตั้งค่าส่วนตัว')
add_image(doc, 'ui-noti.png', 'รูปที่ 12 หน้า Notification')
add_image(doc, 'ui-settings.png', 'รูปที่ 13 หน้า Settings')

# ═══════════════════════════════════════════
# เอกสารอ้างอิง
# ═══════════════════════════════════════════
doc.add_page_break()
add_heading(doc, 'เอกสารอ้างอิง', size=SZ_TITLE, space_before=0)
refs = [
    'Lave, J., & Wenger, E. (1991). Situated Learning: Legitimate Peripheral Participation. Cambridge University Press.',
    'Kolb, D. A. (1984). Experiential Learning: Experience as the Source of Learning and Development. Prentice Hall.',
    'Deterding, S., Dixon, D., Khaled, R., & Nacke, L. (2011). From Game Design Elements to Gamefulness: Defining Gamification. Proceedings of the 15th International Academic MindTrek Conference.',
    'กรมศิลปากร. (2566). ฐานข้อมูลโบราณสถานและประวัติบุคคลสำคัญของไทย. กระทรวงวัฒนธรรม. แหล่งที่มา: https://www.finearts.go.th',
    'Leaflet.js Documentation. (2024). Leaflet — an open-source JavaScript library for mobile-friendly interactive maps [ออนไลน์]. แหล่งที่มา: https://leafletjs.com/reference.html',
    'Supabase. (2024). Supabase Documentation: Auth, Database, and Realtime [ออนไลน์]. แหล่งที่มา: https://supabase.com/docs',
    'OpenStreetMap Contributors. (2024). Thailand Administrative Boundaries GeoJSON [ออนไลน์]. แหล่งที่มา: https://www.openstreetmap.org',
]
for i, ref in enumerate(refs, 1):
    add_para(doc, f'{i}. {ref}', size=SZ_BODY,
             left_indent=0.5, space_after=2, thai_j=True)

# ════════════════════════════════════════════
out_path = r'c:/Users/LENOVO/Documents/Lieutenant Hecker/Primary Data/AI/Coding/Website - NSC prototype 06/document/ตามรอย_NSC_BOOK.docx'
doc.save(out_path)
print(f'Saved: {out_path}')
